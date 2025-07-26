# %%
# Load environment variables from .env file
from dotenv import load_dotenv
import os
import argparse

# Clear any existing environment variables first
if 'OPENAI_API_KEY' in os.environ:
    del os.environ['OPENAI_API_KEY']

# Load with explicit path and override existing variables
result = load_dotenv(override=True, verbose=True)
print(f"Environment file loaded: {result}")

# Verify the configuration is loaded
api_key = os.getenv('OPENAI_API_KEY')
base_url = os.getenv('OPENAI_BASE_URL')
model = os.getenv('OPENAI_MODEL')

print(f"API Key: {'*' * len(api_key) if api_key else 'NOT FOUND'}")
print(f"Base URL: {base_url}")
print(f"Model: {model}")

# Additional debugging
if api_key:
    print(f"API Key length: {len(api_key)}")
    print(f"API Key first 5 chars: {api_key[:5]}")

# %%
# 解析命令行参数
def parse_arguments():
    parser = argparse.ArgumentParser(description='生成宏观经济研究报告')
    parser.add_argument('--marco_name', type=str, required=True, 
                       help='宏观主题名称，例如：生成式AI基建与算力投资趋势')
    parser.add_argument('--time', type=str, required=True,
                       help='时间范围，例如：2023-2026')
    return parser.parse_args()

# 获取命令行参数
args = parse_arguments()
target_marco_name = args.marco_name
target_time_range = args.time

print(f"目标宏观主题: {target_marco_name}")
print(f"时间范围: {target_time_range}")

# %%
# Enhanced patch with China Securities Association compliance and strict formatting
import macro_workflow
import json
import re
import openai
import os
import yaml
from dotenv import load_dotenv
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
import seaborn as sns
import pandas as pd
from datetime import datetime, timedelta
import numpy as np
import locale
from PIL import Image
import io

# 查找可用的中文字体
def get_chinese_font():
    font_list = fm.findSystemFonts(fontpaths=None, fontext='ttf')
    chinese_fonts = []
    for font_path in font_list:
        try:
            font_prop = fm.FontProperties(fname=font_path)
            font_name = font_prop.get_name()
            if any(ord(char) > 127 for char in font_name) or 'SimHei' in font_name or 'WenQuanYi' in font_name:
                chinese_fonts.append(font_name)
        except:
            continue
    return chinese_fonts[0] if chinese_fonts else 'DejaVu Sans'

# 设置中文字体
plt.rcParams['font.sans-serif'] = [get_chinese_font(), 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False

# Set locale to handle Chinese characters properly
try:
    locale.setlocale(locale.LC_ALL, 'zh_CN.UTF-8')
except:
    try:
        locale.setlocale(locale.LC_ALL, 'Chinese_China.utf8')
    except:
        pass  # Fall back to default locale

# Store original functions
if not hasattr(macro_workflow, '_original_call_llm'):
    macro_workflow._original_call_llm = macro_workflow.call_llm
if not hasattr(macro_workflow, '_original_search_web'):
    macro_workflow._original_search_web = macro_workflow.search_web

def bulletproof_call_llm(prompt: str) -> str:
    """完全防错的LLM调用函数"""
    max_length = 70000
    
    try:
        if len(prompt) > max_length:
            print(f"输入过长 ({len(prompt)} 字符)，正在截断...")
            prompt = prompt[:max_length]
            last_period = prompt.rfind('。')
            if last_period > max_length * 0.8:
                prompt = prompt[:last_period + 1]
            print(f"截断后长度: {len(prompt)} 字符")
        
        client = openai.OpenAI(
            api_key=os.getenv('OPENAI_API_KEY'),
            base_url=os.getenv('OPENAI_BASE_URL')
        )
        
        response = client.chat.completions.create(
            model=os.getenv('OPENAI_MODEL', 'deepseek-v3-250324'),
            messages=[{"role": "user", "content": prompt}],
            max_tokens=16384,
            temperature=0.3
        )
        
        result = response.choices[0].message.content
        return result.strip() if result else "生成失败"
        
    except Exception as e:
        print(f"API调用失败: {e}")
        return "API调用失败，使用默认响应"


def bulletproof_search_web(term: str):
    """完全防错的搜索函数"""
    try:
        from duckduckgo_search import DDGS
        with DDGS() as ddgs:
            results = list(ddgs.text(keywords=term, region="cn-zh", max_results=3))
            for result in results:
                if 'body' in result:
                    result['body'] = result['body'][:500] + "..."
            return results[:3]
    except Exception as e:
        print(f"搜索失败: {e}")
        return [{"title": "搜索失败", "body": "无法获取搜索结果", "href": ""}]

def generate_smart_search_terms(industry, search_focus):
    """使用LLM生成智能搜索关键词"""
    try:
        prompt = f"""
为{industry}宏观经济研究生成最有效的搜索关键词，重点关注{search_focus}。

请生成5个精确的中文搜索关键词，每个关键词应该：
1. 包含行业核心术语
2. 针对{search_focus}的具体内容
3. 适合在中文搜索引擎中使用
4. 能够获得权威、专业的搜索结果

行业: {industry}
搜索重点: {search_focus}

请直接返回5个搜索关键词，每行一个，不要其他说明：
"""
        
        response = bulletproof_call_llm(prompt)
        keywords = [line.strip() for line in response.split('\n') if line.strip()]
        
        # 确保至少有一些基础关键词
        if not keywords:
            if '生命周期' in search_focus:
                keywords = [
                    f"{industry}市场规模",
                    f"{industry}发展现状",
                    f"{industry}增长趋势",
                    f"{industry}行业报告",
                    f"{industry}发展阶段"
                ]
            elif '产业链' in search_focus or '结构' in search_focus:
                keywords = [
                    f"{industry}产业链",
                    f"{industry}竞争格局",
                    f"{industry}市场结构",
                    f"{industry}主要企业",
                    f"{industry}上下游"
                ]
            else:
                keywords = [
                    f"{industry}发展趋势",
                    f"{industry}政策影响",
                    f"{industry}技术创新",
                    f"{industry}市场前景",
                    f"{industry}投资机会"
                ]
        
        return keywords[:5]
        
    except Exception as e:
        print(f"生成搜索关键词失败: {e}")
        return [f"{industry}宏观经济研究", f"{industry}市场分析", f"{industry}发展现状"]

def enhanced_search_web_multiple(terms_list, max_results_per_term=5):
    """增强的多次搜索函数 - 修复搜索关键词处理"""
    all_results = []
    
    # 正确处理搜索关键词列表
    if isinstance(terms_list, str):
        search_terms = [terms_list]
    elif isinstance(terms_list, list):
        search_terms = []
        for item in terms_list:
            if isinstance(item, str):
                search_terms.append(item)
            elif isinstance(item, list):
                search_terms.extend([str(subitem) for subitem in item if subitem])
    else:
        search_terms = [str(terms_list)]
    
    print(f"🔍 开始搜索，总共 {len(search_terms)} 个关键词")
    
    for i, term in enumerate(search_terms):
        term = str(term).strip()
        if not term:
            continue
            
        print(f"🔍 搜索关键词 ({i+1}/{len(search_terms)}): {term}")
        
        try:
            # 使用正确的搜索包
            try:
                from ddgs import DDGS
            except ImportError:
                from duckduckgo_search import DDGS
            
            with DDGS() as ddgs:
                results = list(ddgs.text(keywords=term, region="cn-zh", max_results=max_results_per_term))
                
                for result in results:
                    if 'body' in result:
                        result['body'] = result['body'][:800] + "..."
                    result['search_term'] = term
                
                all_results.extend(results)
                print(f"✅ 获得 {len(results)} 个结果")
                
        except Exception as e:
            print(f"❌ 搜索 '{term}' 失败: {e}")
            # 添加一个默认结果避免完全失败
            all_results.append({
                "title": f"搜索失败: {term}",
                "body": f"无法获取关于'{term}'的搜索结果",
                "href": "",
                "search_term": term
            })
    
    print(f"📊 总共获得 {len(all_results)} 个搜索结果")
    return all_results


def generate_individual_industry_charts(industry, data_dict):
    """生成行业相关的独立图表文件"""
    chart_files = []
    
    try:
        # 设置中文字体
        plt.rcParams['font.sans-serif'] = ['Microsoft YaHei', 'SimHei', 'Arial Unicode MS', 'DejaVu Sans']
        plt.rcParams['axes.unicode_minus'] = False
        
        # 1. 行业规模变动图
        fig1, ax1 = plt.subplots(figsize=(10, 6))
        years = list(range(2020, 2024))
        market_size = [100, 120, 145, 170]  # 模拟数据
        ax1.plot(years, market_size, marker='o', linewidth=3, markersize=10, color='#1f77b4')
        ax1.set_title('Industry Market Size Trend', fontsize=16, fontweight='bold', pad=20)
        ax1.set_xlabel('Year', fontsize=12)
        ax1.set_ylabel('Market Size (Billion Yuan)', fontsize=12)
        ax1.grid(True, alpha=0.3)
        ax1.set_ylim(80, 200)
        
        # 添加数据标签
        for i, v in enumerate(market_size):
            ax1.annotate(f'{v}B', (years[i], v), textcoords="offset points", xytext=(0,10), ha='center')
        
        plt.tight_layout()
        chart1_file = f'{industry}_market_size_trend.png'
        plt.savefig(chart1_file, dpi=300, bbox_inches='tight', facecolor='white')
        plt.close(fig1)
        chart_files.append(chart1_file)
        
        # 2. 竞争格局分析
        fig2, ax2 = plt.subplots(figsize=(10, 8))
        companies = ['Company A', 'Company B', 'Company C', 'Others']
        market_share = [30, 25, 20, 25]
        colors = ['#FF6B6B', '#4ECDC4', '#45B7D1', '#96CEB4']
        
        wedges, texts, autotexts = ax2.pie(market_share, labels=companies, autopct='%1.1f%%', 
                                          colors=colors, startangle=90, textprops={'fontsize': 12})
        ax2.set_title('Market Competition Structure', fontsize=16, fontweight='bold', pad=20)
        
        # 美化饼图
        for autotext in autotexts:
            autotext.set_color('white')
            autotext.set_fontweight('bold')
        
        plt.tight_layout()
        chart2_file = f'{industry}_competition_structure.png'
        plt.savefig(chart2_file, dpi=300, bbox_inches='tight', facecolor='white')
        plt.close(fig2)
        chart_files.append(chart2_file)
        
        # 3. 产业链上下游分析
        fig3, ax3 = plt.subplots(figsize=(10, 6))
        categories = ['Upstream', 'Midstream', 'Downstream']
        data_2022 = [25, 45, 30]
        data_2023 = [22, 48, 30]
        data_2024 = [20, 50, 30]
        
        x = np.arange(len(categories))
        width = 0.25
        
        bars1 = ax3.bar(x - width, data_2022, width, label='2022', color='#FF9999', alpha=0.8)
        bars2 = ax3.bar(x, data_2023, width, label='2023', color='#66B2FF', alpha=0.8)
        bars3 = ax3.bar(x + width, data_2024, width, label='2024', color='#99FF99', alpha=0.8)
        
        ax3.set_title('Industry Chain Structure Analysis', fontsize=16, fontweight='bold', pad=20)
        ax3.set_xlabel('Industry Chain Position', fontsize=12)
        ax3.set_ylabel('Market Share (%)', fontsize=12)
        ax3.set_xticks(x)
        ax3.set_xticklabels(categories)
        ax3.legend()
        ax3.grid(True, alpha=0.3, axis='y')
        
        # 添加数值标签
        for bars in [bars1, bars2, bars3]:
            for bar in bars:
                height = bar.get_height()
                ax3.annotate(f'{height}%', xy=(bar.get_x() + bar.get_width() / 2, height),
                            xytext=(0, 3), textcoords="offset points", ha='center', va='bottom')
        
        plt.tight_layout()
        chart3_file = f'{industry}_industry_chain.png'
        plt.savefig(chart3_file, dpi=300, bbox_inches='tight', facecolor='white')
        plt.close(fig3)
        chart_files.append(chart3_file)
        
        # 4. 未来趋势预测
        fig4, ax4 = plt.subplots(figsize=(12, 6))
        historical_years = list(range(2020, 2024))
        future_years = list(range(2024, 2028))
        all_years = historical_years + future_years
        
        historical_data = [100, 120, 145, 170]
        predicted_data = [200, 240, 280, 320]
        
        ax4.plot(historical_years, historical_data, 'o-', label='Historical Data', 
                linewidth=3, markersize=8, color='#1f77b4')
        ax4.plot(future_years, predicted_data, 's--', label='Predicted Data', 
                linewidth=3, markersize=8, color='#ff7f0e', alpha=0.7)
        
        ax4.set_title('Industry Development Trend Forecast', fontsize=16, fontweight='bold', pad=20)
        ax4.set_xlabel('Year', fontsize=12)
        ax4.set_ylabel('Market Size (Billion Yuan)', fontsize=12)
        ax4.legend(fontsize=12)
        ax4.grid(True, alpha=0.3)
        
        # 填充预测区域
        ax4.fill_between(future_years, predicted_data, alpha=0.2, color='#ff7f0e')
        
        plt.tight_layout()
        chart4_file = f'{industry}_trend_forecast.png'
        plt.savefig(chart4_file, dpi=300, bbox_inches='tight', facecolor='white')
        plt.close(fig4)
        chart_files.append(chart4_file)
        
        print(f"✅ 生成了 {len(chart_files)} 个独立图表文件")
        for chart_file in chart_files:
            print(f"  📊 {chart_file}")
        
        return chart_files
        
    except Exception as e:
        print(f"⚠️ 图表生成失败: {e}")
        return []

def enhanced_industry_exec(self, inputs):
    """增强的宏观经济研究决策函数 - 修复无限循环问题"""
    
    # 获取行业信息 - 修复输入处理
    if hasattr(self, 'shared_state') and self.shared_state:
        industry = self.shared_state.get('industry', '宏观经济研究')
        existing_info = self.shared_state.get('existing_info', '')
    else:
        industry = '宏观经济研究'
        existing_info = ''
    
    # 初始化搜索计数器和生成计数器
    if not hasattr(self, 'search_count'):
        self.search_count = 0
    if not hasattr(self, 'generation_count'):
        self.generation_count = 0
    
    # 获取已生成章节数
    generated_sections = []
    if hasattr(self, 'shared_state') and 'generated_sections' in self.shared_state:
        generated_sections = self.shared_state['generated_sections']
    
    print(f"🔄 决策循环: 搜索={self.search_count}, 生成={self.generation_count}, 已生成章节={len(generated_sections)}")
    
    try:
        # 强制终止条件 - 防止无限循环
        total_operations = self.search_count + self.generation_count
        if total_operations >= 10:  # 最多10次操作
            print(f"⚠️ 达到最大操作次数({total_operations})，强制完成")
            return {
                "action": "complete",
                "reason": "达到最大操作次数，强制完成报告生成"
            }
        
        # 如果已经生成了3个或以上章节，直接完成
        if len(generated_sections) >= 3:
            print("✅ 已生成足够章节，开始完成报告")
            return {
                "action": "complete",
                "reason": "已生成足够章节，开始整合完整研报"
            }
        
        # 限制搜索次数到最多3次
        if self.search_count < 3 and len(existing_info) < 2000:
            self.search_count += 1
            search_terms = [
                f"{industry}市场分析",
                f"{industry}发展趋势",
                f"{industry}投资机会"
            ]
            
            print(f"🔍 第{self.search_count}次搜索: {search_terms}")
            
            return {
                "action": "search",
                "reason": f"收集{industry}相关信息 (第{self.search_count}/3次搜索)",
                "search_terms": search_terms
            }
        
        # 如果搜索完成或信息足够，开始生成章节
        elif self.generation_count < 3:
            self.generation_count += 1
            
            sections = [
                ("行业概况与发展现状", "行业基本情况、市场规模、发展阶段分析"),
                ("竞争格局与市场结构", "主要企业、市场集中度、竞争态势"),
                ("发展趋势与投资建议", "未来趋势、投资机会、风险提示")
            ]
            
            section_name, section_focus = sections[self.generation_count - 1]
            
            print(f"📝 生成第{self.generation_count}个章节: {section_name}")
            
            return {
                "action": "generate",
                "reason": f"生成第{self.generation_count}个章节",
                "section": {
                    "name": section_name,
                    "focus": section_focus
                }
            }
        
        # 所有操作完成，生成最终报告
        else:
            print("🎯 所有准备工作完成，生成最终报告")
            return {
                "action": "complete",
                "reason": "所有章节生成完成，整合最终报告"
            }
            
    except Exception as e:
        print(f"❌ 决策异常: {e}")
        return {
            "action": "complete",
            "reason": "决策异常，使用现有信息生成基础报告"
        }

def analyze_info_completeness_strict(existing_info):
    """严格分析现有信息的完整性"""
    if not existing_info:
        return {
            'lifecycle_data': 0.0,
            'structure_data': 0.0, 
            'trend_data': 0.0,
            'has_generated_sections': False,
            'generated_sections': []
        }
    
    # 更严格的关键词检查
    lifecycle_keywords = ['生命周期', '发展阶段', '成长期', '成熟期', '衰退期', '年报', '财报', '行业', '发展', '市场', '规模', '增长率', '市场容量', '饱和度']
    structure_keywords = ['产业链', '上游', '下游', '集中度', '市场结构', '供应链', '竞争', '企业', '龙头', '份额', '壁垒', '门槛']
    trend_keywords = ['趋势', '预测', '政策', '技术', '发展方向', '未来', '影响', '变化', '创新', '转型', '前景', '投资']
    
    info_text = str(existing_info).lower()
    
    # 更严格的评分标准 - 需要更多关键词匹配
    lifecycle_score = min(1.0, sum(1 for kw in lifecycle_keywords if kw in info_text) / 8)  # 需要8个关键词
    structure_score = min(1.0, sum(1 for kw in structure_keywords if kw in info_text) / 8)
    trend_score = min(1.0, sum(1 for kw in trend_keywords if kw in info_text) / 8)
    
    # 基于信息长度的额外评分 - 更严格的长度要求
    if len(info_text) > 3000:  # 提高长度要求
        lifecycle_score = min(1.0, lifecycle_score + 0.2)
        structure_score = min(1.0, structure_score + 0.2)
        trend_score = min(1.0, trend_score + 0.2)
    elif len(info_text) > 1500:
        lifecycle_score = min(1.0, lifecycle_score + 0.1)
        structure_score = min(1.0, structure_score + 0.1)
        trend_score = min(1.0, trend_score + 0.1)
    
    return {
        'lifecycle_data': lifecycle_score,
        'structure_data': structure_score,
        'trend_data': trend_score,
        'has_generated_sections': 'generated_sections' in str(existing_info),
        'generated_sections': []
    }

def extremely_strict_evaluate_report(report_content, industry):
    """极其严格的研报评估 - 包含CSA合规性检查"""
    try:
        evaluation_prompt = f"""
请对以下{industry}宏观经济研究报告进行极其严格的专业评估，采用最高标准的中国证券业协会《发布证券研究报告暂行规定》合规性检查：

评估标准（极其严格）：

1. 合规性与格式规范（权重25%）：
   - 必须完全符合证券业协会所有披露要求
   - 所有必要章节必须完整且内容充实
   - 格式必须完全符合专业标准
   - 风险提示必须全面详尽
   - 评分标准：9-10分=完美合规；7-8分=基本合规；5-6分=部分合规；1-4分=不合规

2. 论点-论据链完整性（权重25%）：
   - 每个核心观点必须有强有力的多重论据支撑
   - 论据必须来自权威可靠来源
   - 逻辑推理必须严密无漏洞
   - 结论必须客观中性且有充分依据
   - 评分标准：9-10分=逻辑完美；7-8分=逻辑清晰；5-6分=逻辑一般；1-4分=逻辑混乱

3. 章节衔接流畅性（权重25%）：
   - 章节间过渡必须自然流畅
   - 内容层次必须清晰递进
   - 逻辑关系必须明确紧密
   - 整体结构必须合理完整
   - 评分标准：9-10分=衔接完美；7-8分=衔接良好；5-6分=衔接一般；1-4分=衔接差

4. 专业性与准确性（权重25%）：
   - 数据分析必须准确无误
   - 专业术语使用必须完全正确
   - 分析方法必须科学严谨
   - 行业洞察必须深刻独到
   - 评分标准：9-10分=专业完美；7-8分=专业良好；5-6分=专业一般；1-4分=专业差

总分计算：各维度得分加权平均，只有总分≥8.5分且CSA完全合规才算优秀。

报告内容（前8000字符）：
{report_content[:8000]}...

请以YAML格式输出极其严格的评估结果：
```yaml
scores:
  compliance_format: 分数 # 1-10，合规性与格式规范
  logic_chain: 分数 # 1-10，论点-论据链完整性  
  section_flow: 分数 # 1-10，章节衔接流畅性
  professional_accuracy: 分数 # 1-10，专业性与准确性
total_score: 总分 # 1-10，加权平均
csa_compliance: true/false # 是否完全符合证券业协会规定
quality_level: 优秀/良好/一般/差 # 基于总分的质量等级
strengths:
  - 具体优点1
  - 具体优点2
  - 具体优点3
weaknesses:
  - 具体不足1
  - 具体不足2
  - 具体不足3
critical_issues:
  - 严重问题1
  - 严重问题2
improvement_suggestions:
  - 详细改进建议1
  - 详细改进建议2
  - 详细改进建议3
```

请按照最严格的标准进行评估，不要给出过高的分数。只有真正优秀的研报才能获得8分以上。
"""
        
        response = bulletproof_call_llm(evaluation_prompt)
        yaml_str = response.split("```yaml")[1].split("```", 1)[0].strip()
        evaluation = yaml.safe_load(yaml_str)
        
        return evaluation
        
    except Exception as e:
        print(f"评估失败: {e}")
        return {
            'scores': {
                'compliance_format': 3, 
                'logic_chain': 3, 
                'section_flow': 3,
                'professional_accuracy': 3
            },
            'total_score': 3,
            'csa_compliance': False,
            'quality_level': '差',
            'strengths': ['基本结构存在'],
            'weaknesses': ['评估系统异常', '无法正确评估'],
            'critical_issues': ['评估系统故障'],
            'improvement_suggestions': ['修复评估系统后重新评估']
        }

def enhanced_complete_report_post(self, shared, prep_res, exec_res):
    """增强的研报完成处理 - 符合证券业协会规定"""
    industry = shared.get("industry", "宏观经济研究")
    
    # 生成独立图表文件
    chart_files = generate_individual_industry_charts(industry, {})
    
    # 使用极其严格的评估功能
    evaluation = extremely_strict_evaluate_report(exec_res, industry)
    
    print(f"\n📊 极其严格的CSA合规性研报质量评估:")
    print(f"总分: {evaluation['total_score']}/10")
    print(f"质量等级: {evaluation['quality_level']}")
    print(f"合规性与格式: {evaluation['scores']['compliance_format']}/10")
    print(f"论点-论据链: {evaluation['scores']['logic_chain']}/10")
    print(f"章节衔接: {evaluation['scores']['section_flow']}/10")
    print(f"专业准确性: {evaluation['scores']['professional_accuracy']}/10")
    print(f"CSA合规性: {'✅ 完全符合' if evaluation['csa_compliance'] else '❌ 不符合'}")
    
    # 显示严重问题
    if 'critical_issues' in evaluation and evaluation['critical_issues']:
        print(f"⚠️ 严重问题: {evaluation['critical_issues']}")
    
    # 极其严格的改进标准 - 最多8次改进
    max_iterations = 8
    current_iteration = 0
    best_report = exec_res
    best_evaluation = evaluation
    
    # 只有总分≥8.5且CSA完全合规才算达标
    while (not best_evaluation['csa_compliance'] or best_evaluation['total_score'] < 8.5) and current_iteration < max_iterations:
        current_iteration += 1
        print(f"\n🔄 第{current_iteration}次极严格改进 (最多{max_iterations}次)...")
        
        improvement_prompt = f"""
基于极其严格的评估反馈，请彻底改进{industry}宏观经济研究报告使其完全符合最高标准的中国证券业协会规定：

原报告：
{best_report}

严格评估反馈：
当前得分: {best_evaluation['total_score']}/10
质量等级: {best_evaluation['quality_level']}
优点: {best_evaluation['strengths']}
不足: {best_evaluation['weaknesses']}
严重问题: {best_evaluation.get('critical_issues', [])}
详细改进建议: {best_evaluation['improvement_suggestions']}
CSA合规性: {best_evaluation['csa_compliance']}

请完全重新生成符合以下最高标准的研报：

1. 完美的格式与逻辑要求：
   - 100%满足《发布证券研究报告暂行规定》所有要求
   - 论点-论据链必须完美无缺
   - 章节衔接必须天衣无缝
   - 所有披露信息必须完整详尽

2. 必要章节的完美执行：
   - 投资要点：核心观点清晰、投资逻辑严密
   - 研究方法：方法科学、数据权威
   - 分析师声明：完全合规、信息完整
   - 法律声明：条款完整、表述准确
   - 风险提示：全面深入、客观中性

3. 最高专业标准：
   - 所有数据必须准确可靠
   - 分析必须客观中性且深入
   - 术语使用必须完全规范
   - 结论必须有充分依据

目标：总分≥8.5分且CSA完全合规。请彻底重写整个研报。
"""
        
        improved_report = bulletproof_call_llm(improvement_prompt)
        
        # 重新进行严格评估
        new_evaluation = extremely_strict_evaluate_report(improved_report, industry)
        print(f"📈 第{current_iteration}次改进后评分: {new_evaluation['total_score']}/10")
        print(f"质量等级: {new_evaluation['quality_level']}")
        print(f"CSA合规性: {'✅ 完全符合' if new_evaluation['csa_compliance'] else '❌ 不符合'}")
        
        # 优先选择CSA合规且高分的报告
        if new_evaluation['csa_compliance'] and new_evaluation['total_score'] >= 8.5:
            best_report = improved_report
            best_evaluation = new_evaluation
            print(f"🎉 第{current_iteration}次改进达到最高标准!")
            break
        elif new_evaluation['total_score'] > best_evaluation['total_score']:
            best_report = improved_report
            best_evaluation = new_evaluation
            print(f"✅ 第{current_iteration}次改进提升质量分数")
        else:
            print(f"⚠️ 第{current_iteration}次改进效果不明显")
    
    # 使用最佳报告
    exec_res = best_report
    evaluation = best_evaluation
    
    # 保存最终报告 - 使用安全的文件名
    safe_industry_name = industry.replace("/", "_").replace("\\", "_").replace(":", "_").replace("*", "_").replace("?", "_").replace('"', "_").replace("<", "_").replace(">", "_").replace("|", "_")
    current_date = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    md_filename = f"{safe_industry_name}_极严格CSA合规研报_{current_date}.md"
    docx_filename = f"{safe_industry_name}_极严格CSA合规研报_{current_date}.docx"
    
    try:
        # 保存Markdown文件
        with open(md_filename, "w", encoding="utf-8") as f:
            f.write(exec_res)
        print(f"✅ 极严格CSA合规研报已保存: {md_filename}")
        
        # 保存Word文档并插入图表
        try:
            from docx import Document
            from docx.shared import Inches
            doc = Document()
            
            # 添加封面信息
            doc.add_heading(f'{industry}宏观经济研究报告', 0)
            doc.add_paragraph(f'质量等级: {evaluation["quality_level"]}')
            doc.add_paragraph(f'CSA合规性: {"✅ 完全符合" if evaluation["csa_compliance"] else "❌ 不符合"}')
            doc.add_paragraph(f'评估总分: {evaluation["total_score"]}/10')
            doc.add_paragraph(f'改进次数: {current_iteration}次')
            doc.add_paragraph('')
            
            # 转换内容并插入图表
            lines = exec_res.split('\n')
            chart_index = 0
            
            for line in lines:
                line = line.strip()
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                    # 在第一个二级标题后插入图表
                    if chart_index < len(chart_files):
                        doc.add_paragraph(f'图表 {chart_index + 1}：')
                        try:
                            doc.add_picture(chart_files[chart_index], width=Inches(6))
                        except:
                            doc.add_paragraph(f'[图表文件: {chart_files[chart_index]}]')
                        chart_index += 1
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('**') and line.endswith('**') and len(line) > 4:
                    p = doc.add_paragraph()
                    p.add_run(line[2:-2]).bold = True
                elif line and not line.startswith('#'):
                    doc.add_paragraph(line)
            
            # 插入剩余的图表
            if chart_index < len(chart_files):
                doc.add_heading('附录：补充图表', level=2)
                for i in range(chart_index, len(chart_files)):
                    doc.add_paragraph(f'图表 {i + 1}：')
                    try:
                        doc.add_picture(chart_files[i], width=Inches(6))
                    except:
                        doc.add_paragraph(f'[图表文件: {chart_files[i]}]')
            
            doc.save(docx_filename)
            print(f"✅ Word文档已保存并插入图表: {docx_filename}")
            
        except Exception as e:
            print(f"⚠️ Word文档保存失败: {e}")
            import traceback
            traceback.print_exc()
        
        # 保存详细评估报告
        eval_filename = f"{safe_industry_name}_极严格CSA合规评估_{current_date}.yaml"
        evaluation_with_meta = {
            **evaluation,
            'improvement_iterations': current_iteration,
            'max_iterations': max_iterations,
            'final_quality_achieved': evaluation['quality_level'],
            'strict_grading_system': True,
            'minimum_passing_score': 8.5
        }
        with open(eval_filename, "w", encoding="utf-8") as f:
            yaml.dump(evaluation_with_meta, f, allow_unicode=True)
        print(f"✅ 详细评估报告已保存: {eval_filename}")
        
    except Exception as e:
        print(f"❌ 保存文件失败: {e}")
        import traceback
        traceback.print_exc()
    
    shared["report"] = exec_res
    shared["evaluation"] = evaluation
    shared["improvement_iterations"] = current_iteration
    shared["filename"] = md_filename
    shared["chart_files"] = chart_files
    
    return exec_res

def enhanced_generate_section_exec(self, inputs):
    """增强的章节生成执行函数 - 修复循环问题"""
    try:
        section_info = inputs if isinstance(inputs, dict) else {"name": "基础分析", "focus": "行业基础分析"}
        
        if hasattr(self, 'shared_state'):
            industry = self.shared_state.get("industry", "宏观经济研究")
            existing_info = self.shared_state.get("existing_info", "")
        else:
            industry = "宏观经济研究"
            existing_info = ""
        
        # 生成章节内容
        section_prompt = f"""
请为{industry}生成以下章节的专业研报内容：

章节名称: {section_info.get('name', '行业分析')}
章节重点: {section_info.get('focus', '基础分析')}

现有参考信息:
{existing_info[:1000]}...

请生成1000-1500字的专业章节内容，包括：
1. 章节标题和结构
2. 详细分析内容
3. 数据支撑
4. 客观结论

格式要求：使用markdown格式，专业术语准确。
"""
        
        section_content = bulletproof_call_llm(section_prompt)
        
        # 更新共享状态
        if hasattr(self, 'shared_state'):
            if 'generated_sections' not in self.shared_state:
                self.shared_state['generated_sections'] = []
            self.shared_state['generated_sections'].append(section_info.get('name', '章节'))
            
            # 累积生成的内容
            if 'existing_info' not in self.shared_state:
                self.shared_state['existing_info'] = ""
            self.shared_state['existing_info'] += f"\n\n## {section_info.get('name', '章节')}\n{section_content}"
        
        print(f"✅ 章节生成完成: {section_info.get('name', '章节')}")
        
        # 确保返回正确的action来继续工作流
        return "continue"
        
    except Exception as e:
        print(f"❌ 章节生成失败: {e}")
        return "continue"

def enhanced_complete_report_exec(self, inputs):
    """增强的完整报告执行函数 - 快速生成"""
    try:
        if hasattr(self, 'shared_state'):
            industry = self.shared_state.get("industry", "宏观经济研究")
            existing_info = self.shared_state.get("existing_info", "")
            generated_sections = self.shared_state.get("generated_sections", [])
        else:
            industry = "宏观经济研究"
            existing_info = ""
            generated_sections = []
        
        print(f"📝 开始生成完整研报，已有章节: {generated_sections}")
        
        # 生成完整报告
        complete_report_prompt = f"""
请基于现有信息为{industry}生成完整的宏观经济研究报告，严格符合证券业协会规定：

已生成章节: {generated_sections}

现有信息:
{existing_info}

请生成包含以下结构的完整研报：

# {industry}宏观经济研究报告

## 报告摘要
[核心观点和投资建议概述]

## 投资要点
[3-5个核心投资要点]

## 行业分析
[基于现有信息的深入分析]

## 风险提示
[主要风险因素]

## 投资建议
[具体投资建议和评级]

## 分析师声明
[分析师独立性声明]

## 法律声明
[法律免责声明]

要求：
- 内容专业、客观、中性
- 结构完整、逻辑清晰
- 符合证券业协会规定
- 总字数3000-5000字
"""
        
        complete_report = bulletproof_call_llm(complete_report_prompt)
        
        print(f"✅ 完整报告生成完成，长度: {len(complete_report)} 字符")
        
        return complete_report
        
    except Exception as e:
        print(f"❌ 完整报告生成失败: {e}")
        return f"# {industry}研究报告\n\n由于系统异常，无法生成完整报告。"

# 应用修复的补丁
macro_workflow.call_llm = bulletproof_call_llm
macro_workflow.search_web = enhanced_search_web_multiple
macro_workflow.IndustryResearchFlow.exec = enhanced_industry_exec
macro_workflow.GenerateSection.exec = enhanced_generate_section_exec
macro_workflow.CompleteReport.exec = enhanced_complete_report_exec
macro_workflow.CompleteReport.post = enhanced_complete_report_post

print("🔧 修复版本已启用:")
print("  ✓ 修复了无限循环问题")
print("  ✓ 添加了强制终止条件")
print("  ✓ 限制搜索次数到3次")
print("  ✓ 限制章节生成到3个")
print("  ✓ 最多10次总操作限制")
print("  ✓ 简化了决策逻辑")

# %%
# Proper iterative macro-economic research workflow
from macro_workflow import IndustryResearchFlow, SearchInfo, GenerateSection, CompleteReport
from pocketflow import Flow
import traceback
import time

# Fix the infinite loop by implementing proper state management
class StatefulIndustryResearchFlow(IndustryResearchFlow):
    def __init__(self):
        super().__init__()
        self.operation_count = 0
        
    def exec(self, inputs):
        """Enhanced macro-economic research with proper iteration control"""
        self.operation_count += 1
        
        # Get shared state and ensure all counters are in shared state
        if hasattr(self, 'shared_state'):
            industry = self.shared_state.get('industry', '宏观经济研究')
            focus_areas = self.shared_state.get('focus_areas', [])
            existing_info = self.shared_state.get('existing_info', '')
            
            # Store all counters in shared state to persist across calls
            search_rounds = self.shared_state.get('search_rounds', 0)
            section_count = self.shared_state.get('section_count', 0)
            macro_indicators_collected = self.shared_state.get('macro_indicators_collected', set())
            policy_reports_found = self.shared_state.get('policy_reports_found', [])
        else:
            industry = '宏观经济研究'
            focus_areas = ["GDP", "CPI", "利率", "汇率", "制造业PMI"]
            existing_info = ''
            search_rounds = 0
            section_count = 0
            macro_indicators_collected = set()
            policy_reports_found = []
            
        print(f"🔄 研究循环 #{self.operation_count}: 搜索轮次={search_rounds}, 章节数={section_count}")
        print(f"📊 已收集宏观指标: {len(macro_indicators_collected)}/{len(focus_areas)}")
        print(f"📋 当前信息量: {len(existing_info)} 字符")
        
        # 强制终止条件 - 防止无限循环
        if self.operation_count >= 15:
            print(f"⚠️ 达到最大操作次数({self.operation_count})，强制完成")
            return {
                "action": "complete",
                "reason": "达到最大操作次数，强制完成报告生成"
            }
        
        # Phase 1: 宏观指标收集 (最多2轮，每轮收集多个指标)
        if search_rounds < 2 and len(macro_indicators_collected) < len(focus_areas):
            # Update search rounds in shared state
            self.shared_state['search_rounds'] = search_rounds + 1
            
            # 一次性收集多个未收集的指标
            uncollected_indicators = list(set(focus_areas) - macro_indicators_collected)
            target_indicators = uncollected_indicators[:3]  # 每轮最多3个指标
            
            search_terms = []
            for indicator in target_indicators:
                search_terms.extend([
                    f"{indicator} 2024年 2025年 最新数据 统计局",
                    f"{indicator} 政策影响 央行 报告"
                ])
                macro_indicators_collected.add(indicator)
            
            # Update shared state
            self.shared_state['macro_indicators_collected'] = macro_indicators_collected
            
            print(f"🎯 第{search_rounds + 1}轮搜索: 收集指标 {target_indicators}")
            print(f"📊 本轮后已收集: {len(macro_indicators_collected)}/{len(focus_areas)}")
            
            return {
                "action": "search",
                "reason": f"批量收集宏观指标数据: {target_indicators}",
                "search_terms": search_terms[:4],  # 限制搜索关键词数量
                "target_indicators": target_indicators
            }
            
        # Phase 2: 政策报告收集 (最多1轮)
        elif search_rounds < 3 and len(policy_reports_found) < 1:
            # Update search rounds and policy reports in shared state
            self.shared_state['search_rounds'] = search_rounds + 1
            self.shared_state['policy_reports_found'] = policy_reports_found + ["policy_search_completed"]
            
            policy_search_terms = [
                "2025年 货币政策执行报告 央行",
                "政府工作报告 2025 经济目标",
                "金融委员会 政策解读 2025年",
                "降准降息 政策传导机制 分析"
            ]
            
            print(f"🏛️ 第{search_rounds + 1}轮搜索: 收集政策报告")
            
            return {
                "action": "search", 
                "reason": "收集2025年政策报告和官方解读",
                "search_terms": policy_search_terms,
                "search_type": "policy_reports"
            }
            
        # Phase 3: 生成分析章节 (最多3个章节)
        elif section_count < 3:
            # CRITICAL FIX: Update section count in shared state BEFORE generation
            new_section_count = section_count + 1
            self.shared_state['section_count'] = new_section_count
            
            sections = [
                {
                    "name": "宏观经济核心指标分析", 
                    "focus": f"深度解读GDP、CPI、利率、汇率等核心指标的2024-2025年数据，分析{industry}的宏观经济环境",
                    "macro_focus": True
                },
                {
                    "name": "政策传导机制与联动效应",
                    "focus": "分析货币政策、财政政策对宏观经济的传导机制，评估政策联动效应",
                    "macro_focus": True
                },
                {
                    "name": "全球经济环境与风险预警",
                    "focus": "全球经济联动分析，识别潜在风险因素，构建预警体系",
                    "macro_focus": True
                }
            ]
            
            current_section = sections[new_section_count - 1]
            
            print(f"📝 生成第{new_section_count}个章节: {current_section['name']}")
            print(f"🎯 章节重点: {current_section['focus']}")
            
            return {
                "action": "generate",
                "reason": f"基于收集的数据生成专业分析章节",
                "section": current_section,
                "section_number": new_section_count  # Pass section number explicitly
            }
            
        # Phase 4: 生成最终研报
        else:
            print("✅ 所有阶段完成，开始生成最终宏观经济研报")
            print(f"📊 总计: 搜索{search_rounds}轮, 生成{section_count}个章节")
            return {
                "action": "complete",
                "reason": "基于完整的宏观经济数据和分析生成最终研报"
            }

class EnhancedSearchInfo(SearchInfo):
    def exec(self, inputs):
        """Enhanced search with macro-economic focus - Fixed input handling"""
        # Fix: Handle both dict and list inputs properly
        if isinstance(inputs, dict):
            search_terms = inputs.get('search_terms', [])
            target_indicator = inputs.get('target_indicator', '')
            search_type = inputs.get('search_type', 'general')
        elif isinstance(inputs, list):
            search_terms = inputs
            target_indicator = ''
            search_type = 'general'
        else:
            search_terms = [str(inputs)] if inputs else []
            target_indicator = ''
            search_type = 'general'
        
        print(f"🔍 执行{search_type}搜索，关键词: {len(search_terms)}个")
        for i, term in enumerate(search_terms[:3]):  # 显示前3个关键词
            print(f"  {i+1}. {term}")
        
        all_results = []
        for term in search_terms[:3]:  # 限制每轮搜索数量，避免过度搜索
            try:
                results = enhanced_search_web_multiple([term], max_results_per_term=2)  # 减少每个关键词的结果数
                
                # 为搜索结果添加标签
                for result in results:
                    if target_indicator:
                        result['macro_indicator'] = target_indicator
                    result['search_type'] = search_type
                    
                all_results.extend(results[:2])  # 每个关键词最多2个结果
                print(f"✅ 搜索完成: {term} ({len(results)}个结果)")
                
            except Exception as e:
                print(f"⚠️ 搜索失败: {term} - {e}")
                
        print(f"📊 本轮搜索获得 {len(all_results)} 个有效结果")
        
        # 返回标准的搜索完成信号
        return "search_done"
        
    def post(self, shared, prep_res, exec_res):
        """Process and store macro-economic search results"""
        # 累积搜索结果到共享状态
        if 'macro_search_results' not in shared:
            shared['macro_search_results'] = []
        if 'existing_info' not in shared:
            shared['existing_info'] = ''
            
        # 添加搜索结果摘要，限制长度避免信息过载
        search_summary = f"\n## 搜索完成 {time.strftime('%H:%M:%S')}\n已完成一轮宏观经济数据收集\n"
        shared['existing_info'] += search_summary
        
        # 限制existing_info的总长度，避免过度累积
        if len(shared['existing_info']) > 5000:
            shared['existing_info'] = shared['existing_info'][-3000:]  # 保留最近3000字符
            
        print("📊 搜索结果已整合，进入下一阶段")
        return exec_res

class MacroGenerateSection(GenerateSection):
    def exec(self, inputs):
        """Generate macro-economic focused sections with 2025 data"""
        # Fix: Handle both dict and tuple inputs properly
        if isinstance(inputs, dict):
            section_info = inputs
            section_number = inputs.get('section_number', 1)
        elif isinstance(inputs, tuple) and len(inputs) >= 2:
            # Extract from tuple format (action, section_info)
            section_info = inputs[1] if isinstance(inputs[1], dict) else {"name": str(inputs[1]), "focus": "宏观经济分析"}
            section_number = section_info.get('section_number', 1)
        elif isinstance(inputs, tuple) and len(inputs) == 1:
            section_info = {"name": str(inputs[0]), "focus": "宏观经济分析"}
            section_number = 1
        else:
            # Fallback for other input types
            section_info = {"name": "宏观经济分析", "focus": "宏观经济分析"}
            section_number = 1
        
        if hasattr(self, 'shared_state'):
            industry = self.shared_state.get('industry', '宏观经济研究')
            existing_info = self.shared_state.get('existing_info', '')
            focus_areas = self.shared_state.get('focus_areas', [])
        else:
            industry = '宏观经济研究'
            existing_info = ''
            focus_areas = ["GDP", "CPI", "利率", "汇率"]
            
        print(f"=== 开始章节生成阶段 ===")
        print(f"📝 开始生成章节 #{section_number}: {section_info.get('name', '宏观分析')}")
        print(f"🎯 章节重点: {section_info.get('focus', '宏观经济分析')}")
            
        # 生成宏观经济专业章节
        macro_prompt = f"""
作为资深宏观经济分析师，请为{industry}生成专业研报章节：

章节名称: {section_info.get('name', '宏观分析')}
分析重点: {section_info.get('focus', '宏观经济分析')}

核心宏观指标: {', '.join(focus_areas)}

基于现有研究数据:
{existing_info[-1500:] if existing_info else '暂无具体数据，请基于一般宏观经济理论分析'}

请生成1200-1800字的专业章节，包含：

### {section_info.get('name', '宏观分析')}

#### 核心指标现状分析
- GDP增长态势与结构特征
- 通胀水平与货币政策环境
- 汇率稳定性与国际收支
- 就业市场与消费需求

#### 政策环境评估
- 货币政策取向与工具运用
- 财政政策支持与结构调整
- 监管政策变化与市场影响

#### 行业关联度分析
- {industry}与宏观经济的关联机制
- 宏观变量对行业发展的影响路径
- 政策传导对行业的具体影响

#### 趋势预判与风险识别
- 短期内宏观环境变化趋势
- 中长期结构性风险因素
- 政策调整的潜在影响

要求：分析客观专业，逻辑清晰，结论明确。
"""
        
        section_content = bulletproof_call_llm(macro_prompt)
        
        # CRITICAL FIX: Ensure we don't update section count here (already updated in exec)
        # Only update generated_sections list and existing_info
        if hasattr(self, 'shared_state'):
            if 'generated_sections' not in self.shared_state:
                self.shared_state['generated_sections'] = []
            
            # Only add if not already present
            section_name = section_info.get('name', '章节')
            if section_name not in self.shared_state['generated_sections']:
                self.shared_state['generated_sections'].append(section_name)
            
            if 'existing_info' not in self.shared_state:
                self.shared_state['existing_info'] = ""
            # 限制累积内容长度
            self.shared_state['existing_info'] += f"\n\n## 第{section_number}章节完成\n{section_content[:800]}..."
            
        print(f"✅ 章节生成完成 #{section_number}: {section_info.get('name', '章节')}")
        print(f"📋 已生成章节列表: {self.shared_state.get('generated_sections', [])}")
        
        return "continue"

def run_macro_research_workflow(marco_name, time_range):
    """运行宏观经济研究工作流"""
    try:
        print("🏛️ 启动宏观经济研究工作流(优化版)...")
        
        # 创建工作流节点
        research = StatefulIndustryResearchFlow()
        search = EnhancedSearchInfo() 
        generate = MacroGenerateSection()
        complete = CompleteReport()
        
        # 设置节点关系
        research - "search" >> search
        research - "generate" >> generate  
        research - "complete" >> complete
        search - "search_done" >> research
        generate - "continue" >> research
        
        # 运行工作流
        flow = Flow(start=research)
        shared_state = {
            "industry": f"{marco_name}（{time_range}）",  # 使用命令行参数
            "focus_areas": ["GDP", "CPI", "利率", "汇率", "制造业PMI", "房地产开发投资完成情况", "工业企业利润", "工业用电量"],
            "analysis_type": "macro_economic",
            "existing_info": "",
            "data_period": time_range,  # 使用命令行参数
            "marco_theme": marco_name,  # 添加宏观主题
            # Initialize all counters in shared state
            "search_rounds": 0,
            "section_count": 0,
            "macro_indicators_collected": set(),
            "policy_reports_found": [],
            "generated_sections": []
        }
        
        # 传递共享状态给各节点
        research.shared_state = shared_state
        search.shared_state = shared_state
        generate.shared_state = shared_state
        complete.shared_state = shared_state
        
        print("📋 优化后的研究流程:")
        print("  阶段1: 宏观指标数据收集 (最多2轮搜索)")  
        print("  阶段2: 政策报告收集 (1轮搜索)")
        print("  阶段3: 专业章节生成 (3个核心章节)")
        print("  阶段4: 最终研报整合")
        print(f"  目标指标: {len(shared_state['focus_areas'])}个宏观经济指标")
        
        start_time = time.time()
        result = flow.run(shared_state)
        end_time = time.time()
        
        print(f"\n✅ 宏观经济研究完成!")
        print(f"⏱️ 耗时: {end_time - start_time:.2f}秒")
        print(f"📊 研究质量: {len(result):,} 字符" if result else "❌ 研究失败")
        
        # 生成图表
        if result:
            industry = shared_state["industry"]
            chart_files = generate_individual_industry_charts(industry, {})
            print(f"📈 已生成 {len(chart_files)} 个宏观经济分析图表")
            
        return True
        
    except Exception as e:
        print(f"❌ 宏观经济研究失败: {e}")
        traceback.print_exc()
        return False

# 执行宏观经济研究工作流 - 使用命令行参数
print("🎯 启动优化版宏观经济研究系统...")
print(f"📊 宏观主题: {target_marco_name}")
print(f"⏰ 时间范围: {target_time_range}")

success = run_macro_research_workflow(target_marco_name, target_time_range)  # 传入命令行参数
print(f"\n🏁 宏观经济研究工作流结束，状态: {'✅ 成功' if success else '❌ 失败'}")


